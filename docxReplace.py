'''
install dependency:
  pip install python-docx

run with:
  python3 script.py

==========================
Script nahradi vsechny slova v dictionary
    ve vsech docx souborech v dane slozce.
Vystup je ve slozce "output".
'''

import os
import re
from docx import Document

dictionary = {"puvodni text": "ZMENENY TEXT", 
              "Testovací input" : "Změněný text"}

ok = True
#  https://stackoverflow.com/questions/24805671/how-to-use-python-docx-to-replace-text-in-a-word-document-and-save
def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def replaceFile(filename):
  doc = Document(filename)
  word_re = (r"puvodni text")
  for word, replacement in dictionary.items():
      word_re=re.compile(word)
      docx_replace_regex(doc, word_re , replacement)

  outputName = 'output/' + filename
  try:
    doc.save(outputName)
    print("Created file: " + outputName)
  except:
    print('\033[91m\tERROR!, ' + outputName + " could not be created, please close this file!\033[0m")


def replaceAll():
  if not os.path.exists('output'):
    os.makedirs('output')
  
  for filename in os.listdir('.'):
    if filename.endswith(".docx") and not filename.startswith("~$"):
      replaceFile(filename)
  
  if ok:
    print('\033[92m\tDone.\033[0m')

replaceAll()